"""다양한 파일 형식에서 텍스트를 추출한다."""


def extract_text(file_obj) -> str:
    """업로드된 파일에서 텍스트를 추출한다.

    지원 형식: txt, pdf, docx, xlsx, xls, csv
    """
    name = file_obj.name.lower()
    data = file_obj.read()

    if name.endswith(".txt") or name.endswith(".csv"):
        return data.decode("utf-8", errors="replace")

    if name.endswith(".pdf"):
        return _read_pdf(data)

    if name.endswith(".docx"):
        return _read_docx(data)

    if name.endswith(".xlsx") or name.endswith(".xls"):
        return _read_excel(data)

    # 알 수 없는 형식 → 텍스트로 시도
    return data.decode("utf-8", errors="replace")


def _read_pdf(data: bytes) -> str:
    try:
        import io
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        return "[PDF 읽기 실패: PyPDF2 패키지가 필요합니다. pip install PyPDF2]"
    except Exception as e:
        return f"[PDF 읽기 오류: {e}]"


def _read_docx(data: bytes) -> str:
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 테이블도 추출
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except ImportError:
        return "[Word 읽기 실패: python-docx 패키지가 필요합니다. pip install python-docx]"
    except Exception as e:
        return f"[Word 읽기 오류: {e}]"


def _read_excel(data: bytes) -> str:
    try:
        import io
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"[Sheet: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)
    except ImportError:
        return "[Excel 읽기 실패: openpyxl 패키지가 필요합니다. pip install openpyxl]"
    except Exception as e:
        return f"[Excel 읽기 오류: {e}]"
