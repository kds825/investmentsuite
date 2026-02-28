"""Report Generator — Markdown / HTML / Word 보고서를 생성한다."""

import os
import re
from datetime import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

TEMPLATE_DIR = Path(__file__).parent / "templates"
OUTPUT_DIR = Path(__file__).resolve().parent.parent.parent / "outputs"


def generate_markdown(context: dict, orchestrator_result: dict, mode: str) -> str:
    """마크다운 형태의 투자보고서를 생성한다."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    company = context.get("company_name", "미정")

    lines = [
        f"# {company} — 투자보고서",
        "",
        f"**모드:** {mode}  ",
        f"**생성 시각:** {now}  ",
        "",
        "---",
        "",
        "## 기본 정보",
        "",
        "| 항목 | 내용 |",
        "|------|------|",
        f"| 업종 | {context.get('industry', '미정')} |",
        f"| 투자 목적 | {context.get('investment_purpose', '미정')} |",
        f"| 투자 기간 | {context.get('investment_period', '미정')} |",
        f"| 엑싯 선호 | {context.get('exit_preference', '미정')} |",
        f"| 리스크 선호 | {context.get('risk_preference', '미정')} |",
        f"| 보고서 톤 | {context.get('report_tone', 'Neutral')} |",
        f"| 체크리스트 깊이 | {context.get('checklist_depth', 'Standard')} |",
        "",
    ]

    # 입력 메모
    memo = context.get("memo", "")
    if memo:
        lines += ["---", "", "## 입력 메모", "", memo, ""]

    # 종합 요약
    summary = orchestrator_result.get("summary", "")
    lines += [
        "---",
        "",
        "## 종합 요약 (Executive Summary & IC-ready Summary)",
        "",
        summary,
        "",
    ]

    # 에이전트 결과
    agent_titles = {
        "industry": "산업 분석 (Industry Analysis)",
        "competitor": "경쟁사 분석 (Competitive Analysis)",
        "consulting": "전략 분석 (Consulting)",
        "dd": "실사 팩 (DD Pack)",
        "legal": "법률 리스크 (Legal)",
        "finance_cost": "재무/비용 분석 (Finance & Cost)",
        "exit_strategy": "엑싯 전략 (Exit Strategy)",
    }

    for key, content in orchestrator_result.get("agent_results", {}).items():
        title = agent_titles.get(key, key)
        lines += ["---", "", f"## {title}", "", content, ""]

    # 푸터
    lines += [
        "---",
        "",
        "> 본 보고서는 LLM 기반 자동 생성 문서입니다. "
        "투자 의사결정 전 반드시 전문가 검토가 필요합니다.",
        f"> 생성 시각: {now} | 모드: {mode}",
    ]

    return "\n".join(lines)


def generate_html(context: dict, orchestrator_result: dict, mode: str) -> str:
    """Jinja2 템플릿으로 HTML 보고서를 렌더링한다."""
    env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)), autoescape=False)
    template = env.get_template("report.html.j2")

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    html = template.render(
        generated_at=now,
        mode=mode,
        company_name=context.get("company_name", "미정"),
        industry=context.get("industry", "미정"),
        investment_purpose=context.get("investment_purpose", "미정"),
        investment_period=context.get("investment_period", "미정"),
        exit_preference=context.get("exit_preference", "미정"),
        risk_preference=context.get("risk_preference", "미정"),
        report_tone=context.get("report_tone", "Neutral"),
        checklist_depth=context.get("checklist_depth", "Standard"),
        memo=context.get("memo", ""),
        summary=orchestrator_result.get("summary", ""),
        agent_results=orchestrator_result.get("agent_results", {}),
    )
    return html


def save_html(html: str, filename: str) -> str:
    """HTML을 outputs/ 폴더에 저장하고 경로를 반환한다."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = OUTPUT_DIR / filename
    filepath.write_text(html, encoding="utf-8")
    return str(filepath)


def save_markdown(md: str, filename: str) -> str:
    """마크다운을 outputs/ 폴더에 저장하고 경로를 반환한다."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = OUTPUT_DIR / filename
    filepath.write_text(md, encoding="utf-8")
    return str(filepath)


def save_docx(md: str, filename: str) -> str | None:
    """마크다운을 Word(.docx)로 변환하여 저장한다. 실패 시 None을 반환."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # PwC 스타일 설정
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        style.paragraph_format.space_after = Pt(4)

        # 헤딩 스타일 — PwC Orange (#D04A02)
        for i in range(1, 5):
            hs = doc.styles[f"Heading {i}"]
            hs.font.color.rgb = RGBColor(0xD0, 0x4A, 0x02)
            hs.font.name = "맑은 고딕"

        for line in md.split("\n"):
            stripped = line.strip()

            # 빈 줄
            if not stripped:
                continue

            # 구분선 — PwC gradient feel
            if stripped == "---":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("━" * 60)
                run.font.color.rgb = RGBColor(0xD0, 0x4A, 0x02)
                run.font.size = Pt(7)
                continue

            # 헤딩
            if stripped.startswith("# "):
                level = 0
                text = stripped.lstrip("# ")
                while stripped[level] == "#":
                    level += 1
                text = stripped[level:].strip()
                heading = doc.add_heading(text, level=min(level, 4))
                continue

            # 테이블 구분선 (|---|---| 등) 스킵
            if stripped.startswith("|") and set(stripped.replace("|", "").replace("-", "").strip()) == set():
                continue

            # 테이블 헤더/행 감지 → 연속된 테이블 행을 모아서 처리
            if stripped.startswith("|") and stripped.endswith("|"):
                # 테이블 행을 파싱
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                # 단일 행을 paragraph로 표현 (docx 테이블은 행 수를 미리 알아야 하므로)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                for i, cell in enumerate(cells):
                    if i > 0:
                        p.add_run("  |  ")
                    run = p.add_run(cell)
                    # 볼드 처리 (**text**)
                    clean = cell.replace("**", "")
                    if "**" in cell:
                        run.text = clean
                        run.bold = True
                    run.font.size = Pt(9)
                continue

            # 인용 (blockquote) — PwC Amber accent
            if stripped.startswith("> "):
                text = stripped[2:]
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run("│ ")
                run.font.color.rgb = RGBColor(0xEB, 0x8C, 0x00)
                run.bold = True
                run = p.add_run(text)
                run.italic = True
                run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
                continue

            # 불릿 리스트
            if stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:]
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_run(p, text)
                continue

            # 번호 리스트
            if re.match(r"^\d+[\.\)] ", stripped):
                text = re.sub(r"^\d+[\.\)] ", "", stripped)
                p = doc.add_paragraph(style="List Number")
                _add_formatted_run(p, text)
                continue

            # 볼드 전용 줄 (**text**)
            if stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(stripped.strip("* "))
                run.bold = True
                continue

            # 일반 텍스트
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        filepath = OUTPUT_DIR / filename
        doc.save(str(filepath))
        return str(filepath)

    except Exception:
        return None


def _add_formatted_run(paragraph, text: str):
    """볼드(**text**) 마크다운을 파싱하여 run에 추가한다."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
